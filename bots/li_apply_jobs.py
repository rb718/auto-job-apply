import time, random, os, csv, platform
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.keys import Keys
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.select import Select
from bs4 import BeautifulSoup
import pandas as pd
import pyautogui
import logging
from urllib.request import urlopen
from webdriver_manager.chrome import ChromeDriverManager
import re
import yaml
import os
from datetime import datetime, timedelta
import json

driver = webdriver.Chrome(ChromeDriverManager().install())
wkdir = os.getcwd()

class EasyApplyBot:

    MAX_SEARCH_TIME = 10*600

    def __init__(self,fullName,username,password,max_jobs_to_apply,uploads={},filename='output.csv',blacklist=[]):

        logging.info("\nWelcome to Easy Apply Bot\n")
        import os
        import datetime

        emailUsr = (username.split('@'))[0]
        self.full_name = fullName
        self.user = emailUsr
        self.max_jobs_to_apply = max_jobs_to_apply
        self.dirpath = os.getcwd()
        self.uploads = uploads
        past_ids = self.get_appliedIDs(filename)
        self.appliedJobIDs = past_ids if past_ids != None else []
        self.filename = filename
        self.options = self.browser_options()
        self.browser = driver
        self.wait = WebDriverWait(self.browser, 30)
        self.blacklist = blacklist
        self.start_linkedin(username, password)
        TODAY = datetime.datetime.now()
        self.TODAY = TODAY.strftime("%B-%d-%Y")

    def get_appliedIDs(self, filename):
        try:
            path_to_current = os.path.join(wkdir, filename)
            if os.path.isfile(path_to_current) is False:
                with open(filename,'w+',newline='') as opFile:
                    writer = csv.writer(opFile)
                    writer.writerow(["Date", "Job ID", "Job Title","Company Name","Attemp Count","Job Listing Date","Number of Applicants","Job Description","Applied"])
            
            df = pd.read_csv(filename,
                            header=0,
                            names=['timestamp', 'jobID', 'job', 'company', 'attempted','jld','noa','jd', 'result'],
                            lineterminator='\n',
                            encoding='latin1',
                            verbose = True,
                            skip_blank_lines=True)
            # df['timestamp'] = pd.to_datetime(df['timestamp'], format="%Y-%m-%d %H:%M:%S.%f")
            # df = df[df['timestamp'] > (datetime.now() - timedelta(days=2))]
            df_ = df.loc[df.result == 'TRUE\r']
            jobIDs = list(df_.jobID)
            logging.info(f"{len(jobIDs)} jobIDs found")
            return jobIDs
        except Exception as e:
            logging.exception(str(e) + "   jobIDs could not be loaded from CSV {}".format(filename))
            return None

    def browser_options(self):
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--ignore-certificate-errors")
        options.add_argument('--no-sandbox')
        options.add_argument("--disable-extensions")
        return options

    def start_linkedin(self,username,password):
        logging.info("\nLogging in.....\n \nPlease wait :) \n ")
        self.browser.get("https://www.linkedin.com/login?trk=guest_homepage-basic_nav-header-signin")
        try:
            user_field = self.browser.find_element_by_id("username")
            user_field.clear()
            pw_field = self.browser.find_element_by_id("password")
            pw_field.clear()
            login_button = self.browser.find_element_by_css_selector(".btn__primary--large")
            # user_field.send_keys(username)
            for char in username:
                user_field.send_keys(char)
                randomSec = random.uniform(0.0051, 0.01)
                time.sleep(randomSec)

            user_field.send_keys(Keys.TAB)
            time.sleep(1)
            # pw_field.send_keys(password)
            for char in password:
                pw_field.send_keys(char)
                randomSec = random.uniform(0.0051, 0.01)
                time.sleep(randomSec)
            time.sleep(1)
            login_button.click()

            crntUrl = self.browser.current_url
            retryCount =0
            while ('/checkpoint/challenge/' in crntUrl) and (retryCount < 5):
                logging.critical('Login Failed! Solve login challenge', retryCount)
                pin_verify = self.browser.find_element_by_id('input__email_verification_pin')
                if pin_verify:
                    logging.critical('\n\n\n!!!!!! Got Email Pin Code, Waiting for it to be resolved. !!!!!!\n\n\n')
                time.sleep(5)
                retryCount += retryCount + 1

        except TimeoutException:
            logging.exception("\n\n\n!!!!!!TimeoutException! Username/password field or login button not found !!!!!!\n\n\n")

    def fill_data(self):
        self.browser.set_window_size(0, 0)
        self.browser.set_window_position(2000, 2000)

    # ---ENTRYPOINT--- #
    def start_apply(self, positions, locations):
        start = time.time()
        self.fill_data()

        combos = []
        while len(combos) < len(positions) * len(locations):
            position = positions[random.randint(0, len(positions) - 1)]
            location = locations[random.randint(0, len(locations) - 1)]
            combo = (position, location)
            if combo not in combos:
                combos.append(combo)
                logging.info(f"Applying to {position}: {location}")
                location = "&location=" + location
                apldJobs = self.applications_loop(position, location,[],None)
            if len(combos) > 20:
                break
        self.finish_apply()

    def applications_loop(self, position, location,customJobIds,customUri):
        count_application = 0
        count_job = 0
        jobs_per_page = 0
        start_time = time.time()
        max_jobs_applied = False
        all_job_ids = []
        jobListPageJobIds=None
        logging.info("\nLooking for jobs.. Please wait..")

        self.browser.set_window_position(0, 0)
        self.browser.maximize_window()
        
        if (len(customJobIds) == 0) and (customUri is None):
            logging.info('\t Search Type: Automatic')
            self.browser, jobs_per_page,lstPageIds = self.next_jobs_page(position, location, jobs_per_page)
            jobListPageJobIds = None
            jobListPageJobIds = lstPageIds
        elif customUri:
            logging.info('\t Search Type: Custom URI')
            self.browser,lstPageIds = self.custom_jobs_page(customUri)
            jobListPageJobIds = None
            jobListPageJobIds = lstPageIds
        elif len(customJobIds) >0:
            logging.info('\t Search Type: Custom JOBID')
        else:
            logging.critical('\t Search Type: __MISSING__')
        
        logging.info("Looking for jobs.. Please wait..\n")

        while time.time() - start_time < self.MAX_SEARCH_TIME:
            logging.info(f"{(self.MAX_SEARCH_TIME - (time.time() - start_time))//60} minutes left in this search")

            # sleep to make sure everything loads, add random to make us look human.
            time.sleep(random.uniform(3.5, 6.9))
            self.load_page(sleep=1)

            if len(customJobIds) == 0:

                # get job Ids
                IDs = jobListPageJobIds

                # remove already applied jobs
                before = len(IDs)
                jobIDs = [x for x in IDs if x not in self.appliedJobIDs]
                after = len(jobIDs)

                if len(jobIDs) == 0 and len(IDs) > 24:
                    jobs_per_page = jobs_per_page + 25
                    count_job = 0
                    self.avoid_lock()
                    # Handle below for custom url
                    if (len(customJobIds) == 0) and (customUri is None):
                        logging.info('\t Search Type: Automatic')
                        self.browser, jobs_per_page,lstPageIds = self.next_jobs_page(position, location, jobs_per_page)
                        jobListPageJobIds = None
                        jobListPageJobIds = lstPageIds
                    elif customUri:
                        logging.info('\t Search Type: Custom URI')
                        self.browser, jobs_per_page, lstPageIds = self.next_jobs_page_custom_uri(customUri,jobs_per_page)
                        jobListPageJobIds = None
                        jobListPageJobIds = lstPageIds

                    # self.browser, jobs_per_page = self.next_jobs_page(position,
                    #                                                 location,
                    #                                                 jobs_per_page)
            else:
                jobIDs = customJobIds
            
            logging.info('\n**** MAX JOBS TO APPLY FOR:{a} *****'.format(a=self.max_jobs_to_apply))
            logging.info(f'\nExtracted JOB ID"s from Job Listing Page: {jobIDs}')

            # loop over IDs to apply
            for i, jobID in enumerate(jobIDs):
                logging.info(f"\n\n***********************NEW APPLICATION STARTED ***********************")
                
                try:
                    continueJob = True
                    obj = self.readJobConfig()
                    continueJob = obj.get('run')
                    logging.info(f'\t\t Continue Next Job or Move:{continueJob}')
                    while not continueJob:
                        logging.warning(f"Bot is in {continueJob} state. as User has passed Pause status. Sleeping for 20 sec")
                        time.sleep(20)
                        obj = self.readJobConfig()
                        continueJob = obj.get('run')

                except Exception as e:
                    logging.exception(f"{e}")

                all_job_ids.append(jobID)
                logging.info('\t\tCurrent Job Page Count: {i}'.format(i=(i+1)))
                logging.info('\t\tCount of Jobs Applied:{a}'.format(a=count_application))
                logging.info('\t\tNumber of TOTAL Jobs For Current listing URL:{a}'.format(a=len(jobIDs)))
                logging.info(f'\t\tCurrent Job URL is https://www.linkedin.com/jobs/view/{jobID}/? \n')
                self.get_job_page(jobID)

                if i > 0:
                    #Random Wait Between Job Application
                    randomSec = random.uniform(321.6, 624.6)
                    logging.info('WAIT TIME:{randomSec} | Bot is in wait state. It will auto-start'.format(randomSec=randomSec))
                    time.sleep(randomSec)

                jd={}
                try:
                    jd['job_id']=jobID

                    jd_post = self.browser.find_element_by_xpath('//section[@class="artdeco-card ember-view"]/div[2]//p')
                    if jd_post:
                        jd_posted = jd_post.text
                        jd_posted_text = jd_post.get_attribute("innerText")
                    reMatch = re.search('Posted Date(.|\n*)Posted (.*)(\n*)Number of applicants(\n)(.*)applicants',jd_posted, re.IGNORECASE)
                    if reMatch.group(2):
                        jd['date_listed']=reMatch.group(2)
                    else:
                        jd['date_listed']='NA'
                    if reMatch.group(5):
                        jd['number_of_applicants']=reMatch.group(5)
                    else:
                        jd['number_of_applicants']='NA'

                    job_descr=self.browser.find_element_by_xpath('//div[@id="job-details"]//span')
                    if job_descr:
                        job_description = job_descr.text
                        job_description_text = job_descr.get_attribute("innerText")
                    
                    jd['job_description']=job_description_text.replace(',','_')
                    jd['job_description_raw']=job_description.replace(',','_')
                    jd['job_title'] = self.browser.find_element_by_xpath("//div[contains(@class,'flex-grow-1')]//h1").text
                    jd['job_company'] = self.browser.find_element_by_xpath("//h3[contains(@class,'jobs-top-card__company-info')]//a[contains(@data-control-name,'company_link')]").text
                    jd['job_location'] = self.browser.find_element_by_xpath("//h3[contains(@class,'jobs-top-card__company-info')]//span[2]").text
                    jd['job_country'] = self.browser.find_element_by_xpath("//h3[contains(@class,'jobs-top-card__company-info')]//span[4]").text
                except:
                    pass
                # get easy apply button
                button = self.get_easy_apply_button()
                if button is not False:
                    string_easy = "* has Easy Apply Button"
                    time.sleep(0.7)
                    button.click()
                    time.sleep (2.8)
                    result,jd = self.send_resume(jd)
                    
                    count_application += 1
                    count_job += 1
                else:
                    string_easy = "* Doesn't have Easy Apply Button"
                    result = False

                position_number = str(count_job + jobs_per_page)
                logging.info(f"\nPosition {position_number}:\n {self.browser.title} \n {string_easy} \n")
                try:
                    self.write_to_file(button, jobID, self.browser.title, result,jd)
                except Exception as e:
                    logging.exception(str(e))
                    pass
                
                # Break if today's max jobs are applied
                if count_application == self.max_jobs_to_apply:
                    logging.info(':::::::::::::::::::::::::: MAX JOBS SUCCESSFULLY APPLIED ::::::::::::::::::::::::::')
                    max_jobs_applied=True
                    break
                # sleep every 20 applications
                if count_application != 0  and count_application % 20 == 0:
                    sleepTime = random.randint(200, 300)
                    
                    logging.info(f'\n********count_application: {count_application}************\n\n')
                    logging.info(f'Multiple of 20 Job already applied. Bot is in wait state. It will auto-start')
                    logging.info(f"Time for a nap - see you in:{int(sleepTime/60)} min\n")
                    time.sleep(sleepTime)

                # go to new page if all jobs are done
                if count_job == len(jobIDs) - 1 and len(customJobIds) == 0:
                    jobs_per_page = jobs_per_page + 25
                    count_job = 0

                    logging.info('\n\n****************************************\n\n')
                    logging.info('Going to next jobs page, YEAAAHHH!!')
                    logging.info('\n\n****************************************\n\n')

                    self.avoid_lock()
                    if customUri is None:
                        logging.info('_Automatic_Search_Next_URL_')
                        self.browser, jobs_per_page, lstPageIds = self.next_jobs_page(position,location,jobs_per_page)
                        jobListPageJobIds = None
                        jobListPageJobIds = lstPageIds
                    elif customUri:
                        logging.info('_Custom_Search_Next_URL_')
                        self.browser, jobs_per_page, lstPageIds = self.next_jobs_page_custom_uri(customUri,jobs_per_page)
                        jobListPageJobIds = None
                        jobListPageJobIds = lstPageIds
                    else:
                        logging.info('No NEXT PAGE. Probably, I"ts Custom JobID viz. singular ')
                        pass
            if len(jobIDs) == 0:
                break
            if max_jobs_applied:
                break
        return all_job_ids
    
    def readJobConfig(self):
        obj ={}
        try:
            with open('threadConfig.json','r') as f:
                obj = json.loads(f.read())
        except Exception as e:
            logging.exception(f"{e}")
        return obj
    
    def write_to_file(self, button, jobID, browserTitle, result,metadata):
        try:
            path_to_current = os.path.join(wkdir, self.filename)
            if os.path.isfile(path_to_current) is False:
                with open(self.filename,'w+',newline='') as opFile:
                    writer = csv.writer(opFile)
                    writer.writerow(["Date", "Job ID", "Job Title","Company Name","Attemp Count","Job Listing Date","Number of Applicants","Job Description","Applied"])

            def re_extract(text, pattern):
                target = re.search(pattern, text)
                if target:
                    target = target.group(1)
                return target

            timestamp = datetime.now()
            attempted = False if button == False else True
            job = re_extract(browserTitle.split(' | ')[0], r"\(?\d?\)?\s?(\w.*)")
            company = re_extract(browserTitle.split(' | ')[1], r"(\w.*)" )

            toWrite = [timestamp, jobID, job, company, attempted, metadata['date_listed'],metadata['number_of_applicants'],(metadata['job_description_raw']).replace("\u202f","_"),result]
            with open(self.filename,'a',newline='') as f:
                writer = csv.writer(f)
                writer.writerow(toWrite)
        except Exception as exc:
            logging.exception(str(exc))
            pass

    def get_job_page(self, jobID):

        job = 'https://www.linkedin.com/jobs/view/'+ str(jobID)
        self.browser.get(job)
        self.job_page = self.load_page(sleep=0.5)
        return self.job_page

    def get_easy_apply_button(self):
        try :
            button = self.browser.find_elements_by_xpath(
                        '//button[contains(@class, "jobs-apply")]/span[1]'
                        )

            EasyApplyButton = button[0]
        except :
            EasyApplyButton = False

        return EasyApplyButton

    def send_resume(self,metadata={}):
        def is_present(button_locator):
            return len(self.browser.find_elements(button_locator[0],
                                                     button_locator[1])) > 0
        try:
            time.sleep(random.uniform(1.5, 2.5))
            next_locater = (By.CSS_SELECTOR,"button[aria-label='Continue to next step']")
            review_locater = (By.CSS_SELECTOR,"button[aria-label='Review your application']")
            submit_locater = (By.CSS_SELECTOR,"button[aria-label='Submit application']")
            submit_application_locator = (By.CSS_SELECTOR,"button[aria-label='Submit application']")
            error_locator_OLD = (By.CSS_SELECTOR,"p[data-test-form-element-error-message='true']")
            error_locator_radio = (By.XPATH,"//fieldset[@aria-invalid='true']//p[@data-test-form-element-error-message='true']")
            error_locator_sl = (By.XPATH,"//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']")
            upload_locator = (By.CSS_SELECTOR,"input[name='file']")
            deselect_used_resume = (By.CSS_SELECTOR,"button[aria-label='Remove uploaded document']")
            form_questions = (By.XPATH,"//form/div[@class='jobs-easy-apply-form__groupings']//label")
            form_questions_title = (By.XPATH,"//h3")
            form_questions_radio = (By.XPATH,"//div[@class='jobs-easy-apply-form-section__grouping']//legend/span[1]")
            resume_download_button= (By.XPATH,"//button[@aria-label='Download uploaded document']")
            resume_upload_retry = (By.XPATH,"//button[@aria-label='Click to try again']")
            follow_company = (By.XPATH,"//input[@id='follow-company-checkbox']")
            follow_company_attr = (By.XPATH,"//input[@id='follow-company-checkbox']//../label")
            terms_and_condition = (By.XPATH,"//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//../div//label")

            # labels
            qna_labels = (By.XPATH,"//label[contains(@class,'fb-form-element-label')]")
            qna_legends = (By.XPATH,"//legend[contains(@class,'fb-form-element-label')]")
            cnclWE = (By.XPATH,"//button[@data-control-name='cancel_edit_unify']")

            submitted = False
            questions = []
            while True:
                radioDone = True
                slDone = True
                countResumeUpload = 0
                # Upload Cover Letter if possible
                
                if is_present(deselect_used_resume) or is_present(upload_locator):
                    logging.info('\t *File Upload Field Detected.')
                    if is_present(deselect_used_resume) > 0:
                        logging.info('\t *Existing Deselect Button Detected.')
                        self.browser.find_element(deselect_used_resume[0],deselect_used_resume[1]).click()
                    if is_present(upload_locator):
                        logging.info('\t *Uploader Input Box Is Present')
                        input_buttons = self.browser.find_elements(upload_locator[0],
                                                                upload_locator[1])
                        for input_button in input_buttons:
                            parent = input_button.find_element(By.XPATH, "..")
                            sibling = parent.find_element(By.XPATH, "preceding-sibling::*")
                            grandparent = sibling.find_element(By.XPATH, "..")

                            uploads = {'Resume': '/path/to_resume'}
                            for key in uploads.keys():
                                if key == 'Resume':
                                    logging.info(f'\t *UPLOADS_KEY_CHECK{key}')
                                    file_name = str(metadata['job_id'])+'_'+(self.full_name).replace(' ','_')
                                    metadata['file_name']=file_name
                                    
                                    # Create Resume to Keep in Today's day folder
                                    self.create_resume(metadata)

                                    if (key in sibling.text) or (key in grandparent.text) or ('Upload resume' in parent.text) or ('Upload resume' in sibling.text):
                                        try:
                                            logging.info(f'\t *Filename Is: {file_name}')

                                            filepath = os.path.join(wkdir,'resume\\'+str(self.TODAY)+'\\'+file_name+'.pdf')
                                            input_button.send_keys(filepath)
                                            countResumeUpload = countResumeUpload + 1

                                            logging.warning('\t\t *reCHECK-cross-button-state: {i}'.format(i=is_present(deselect_used_resume)))
                                            while is_present(resume_download_button) == 0:
                                                logging.info('\t\t *Waiting for the resume to get uploaded!!!')
                                                time.sleep(1)
                                                if is_present(resume_upload_retry) >0 :
                                                    self.browser.find_element(resume_upload_retry[0],resume_upload_retry[1]).click()
                                        except Exception as fileUploadException:
                                            logging.exception(str(fileUploadException))
                                            break
                                
                    #input_button[0].send_keys(self.cover_letter_loctn)
                    time.sleep(random.uniform(4.5, 6.5))

                # Click Next or submitt button if possible
                button = None
                buttons = [next_locater,review_locater,submit_locater] # 
                for i, button_locator in enumerate(buttons):
                    logging.info(f'\n\t || ITERATEING BTNS: {i} ||')
                    
                    try:
                        '''
                        #. Check if File Upload Button is pressed
                        #. Check is button is present
                        #. Check is qna are present
                        #. Check t&C button is present
                        #. Check if company Follow button is present
                        #. Click Button

                        '''
                        if (is_present(deselect_used_resume) or is_present(upload_locator)) and countResumeUpload == 0:
                            logging.warning('\t\t**DETECTED: Secondary RESUME')
                            if is_present(deselect_used_resume) > 0:
                                logging.info('\t\t**Deselected button')
                                self.browser.find_element(deselect_used_resume[0],deselect_used_resume[1]).click()
                            if is_present(upload_locator):
                                logging.info('\t\t**Uploader Input Box Is present')
                                input_buttons = self.browser.find_elements(upload_locator[0],
                                                                        upload_locator[1])
                                for input_button in input_buttons:
                                    parent = input_button.find_element(By.XPATH, "..")
                                    sibling = parent.find_element(By.XPATH, "preceding-sibling::*")
                                    grandparent = sibling.find_element(By.XPATH, "..")

                                    uploads = {'Resume': '/path/to_resume'}
                                    for key in uploads.keys():
                                        if key == 'Resume':
                                            logging.info(f'\t\t**UPLOADS_KEY_CHECK: {key}')
                                            file_name = str(metadata['job_id'])+'_'+(self.full_name).replace(' ','_')
                                            metadata['file_name']=file_name
                                            
                                            # Create Resume to Keep in Today's day folder
                                            self.create_resume(metadata)

                                            if (key in sibling.text) or (key in grandparent.text) or ('Upload resume' in parent.text) or ('Upload resume' in sibling.text):
                                                try:
                                                    logging.info(f'\t\t**FILENAME IS:{file_name}')
                                                    filepath = os.path.join(wkdir,'resume\\'+str(self.TODAY)+'\\'+file_name+'.pdf')
                                                    input_button.send_keys(filepath)
                                                    countResumeUpload = countResumeUpload + 1
                                                    logging.info('\t\t**reCHECK-cross-button-state{i}'.format(i=is_present(deselect_used_resume)))
                                                    while is_present(resume_download_button) == 0:
                                                        logging.info('\t\t**waiting for the resume to get uploaded!!!')
                                                        time.sleep(1)
                                                        if is_present(resume_upload_retry) >0 :
                                                            self.browser.find_element(resume_upload_retry[0],resume_upload_retry[1]).click()
                                                except Exception as fileUploadException:
                                                    logging.exception(str(fileUploadException))
                                                    break
                                    
                            #input_button[0].send_keys(self.cover_letter_loctn)
                            time.sleep(random.uniform(4.5, 6.5))
                        if is_present(button_locator):
                            logging.info('          Button is present. Checking its clickable state')
                            logging.info(f'          {button_locator}')
                            button = self.wait.until(EC.element_to_be_clickable(button_locator))
                        if is_present(error_locator_radio) or is_present(error_locator_sl):
                            logging.critical('\n**Error Found Due to Unanswered Questions.**\n')
                            button = None
                            break
                        if is_present(qna_labels) or is_present(qna_legends):
                            isReqd = False
                            
                            qnaLoopResult = {}
                            if is_present(qna_labels):
                                labels = self.browser.find_elements(qna_labels[0],qna_labels[1])
                                qnaLoopLabelsResult = self.loop_qna(labels)
                                qnaLoopResult = self.Merge(qnaLoopResult,qnaLoopLabelsResult)
                            if is_present(qna_legends):
                                labels = self.browser.find_elements(qna_legends[0],qna_legends[1])
                                qnaLoopLegendsResult = self.loop_qna(labels)
                                qnaLoopResult = self.Merge(qnaLoopResult,qnaLoopLegendsResult)
                            
                            # check mandates
                            logging.info(f"{json.dumps(qnaLoopResult)}")
                            slDone = qnaLoopResult['slDone']
                            radioDone = qnaLoopResult['radioDone']
                        # Other elements --- Sign Terms & Conditions, Unfollow company
                        if (radioDone is False) or (slDone is False):
                            logging.critical("Could not complete submission | slDone or radioDone unanswered")
                            break
                        if is_present(terms_and_condition):
                            checked = self.browser.find_element_by_xpath("//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//../div//label/../input").get_attribute('checked')
                            if not checked:
                                self.browser.find_element_by_xpath("//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//../div//label").click()
                        if is_present(follow_company):
                            checked = self.browser.find_element(follow_company[0],follow_company[1]).get_attribute('checked')
                            if checked:
                                logging.info('          Unchecking Follow company')
                                self.browser.find_element(follow_company_attr[0],follow_company_attr[1]).click()
                        if button:
                            logging.info(f'          No SL, Radio, DL NOW - Pressing Button with index: {i}')
                            button.click()
                            time.sleep(random.uniform(1.5, 2.5)) #TODO randomize sleep here
                            if (i == 2) and (not is_present(button_locator)):
                                logging.info('          Submit  Button Pressed')
                                submitted = True
                            elif (i == 2) and is_present(button_locator):
                                button.click()
                                if (i == 2) and (not is_present(button_locator)):
                                    logging.info('          Submit  Button Pressed by retry')
                                    submitted = True
                            break
                        else:
                            logging.info('          No SL, Radio, DL or Button')
                    except Exception as e:
                        logging.exception(str(e))
                        self.writeError(str(e))
                        button = None
                        pass
                
                if button == None:
                    logging.critical("Could not complete submission | Button is None")
                    break
                elif (radioDone is False) or (slDone is False):
                    logging.critical("Could not complete submission | slDone or radioDone unanswered")
                    break
                elif submitted:
                    logging.info("Application Submitted")
                    break
            
            time.sleep(random.uniform(1.5, 2.5))

            #After submiting the application, a dialog shows up, we need to close this dialog
            close_button_locator = (By.XPATH, "//button[(contains(@aria-label,'Dismiss')) and not (contains(@type,'button'))]")
            if is_present(close_button_locator):
                try:
                    close_button = self.wait.until(EC.element_to_be_clickable(close_button_locator))
                    close_button.click()
                except Exception as e:
                    logging.exception(str(e))
                    pass

        except Exception as e:
            logging.exception(str(e))
            logging.warning("cannot apply to this job")
            raise(e)

        return submitted,metadata

    def loop_qna(self,labels):
        logging.info('\t\t__________ENTERED_______\t\t')
        metaInfo ={}
        quesObj ={}
        text =''
        slDone=True
        radioDone = True
        with open('questions.json') as jf:
            quesObj = json.load(jf)

        for label in labels:
            labelTxt = None
            elem_label = None
            elem_legend = None
            
            try:
                elem_label = label.find_element(By.XPATH,'../label')
            except Exception as lblExcept:
                pass
            try:
                elem_legend = label.find_element(By.XPATH,'../legend')
            except Exception as lgndExcept:
                pass

            if elem_label:
                metaInfo['elem_type']='LABEL'
                # logging.info(f"It's label type")
                labelTxt = (label.find_element(By.XPATH,'../label').text).strip()
            elif elem_legend:
                metaInfo['elem_type']='LEGEND'
                # logging.info(f"It's legend type")
                labelTxt = (label.find_element(By.XPATH,'../legend').text).strip()
            
            actualLabelText = labelTxt.split('\n')[0]
            # logging.info(f"Label Text Extracted is : {actualLabelText}")
            metaInfo[actualLabelText] ={}
            # Check it's Required Field or Not
            if 'Required' not in labelTxt:
                isReqd = False
            else:
                isReqd = True
            metaInfo[actualLabelText]['isRequired'] =isReqd
            # Check Type of field
            fieldTypeLocator = None
            fieldType = None
            fieldSatisfied = False
            getFieldValue =None

            if len(label.find_elements(By.XPATH,"..//input")):
                fieldTypeLocator = label.find_element(By.XPATH,"..//input")
                if fieldTypeLocator:
                    fieldType = fieldTypeLocator.get_attribute('type')
                    getFieldValue = fieldTypeLocator.get_attribute('value').strip()
                    if getFieldValue and (getFieldValue != '1') and (getFieldValue != 1):
                        fieldSatisfied = True
                    else:
                        fieldSatisfied = False
            elif len(label.find_elements(By.XPATH,"..//select")):
                fieldTypeLocator = label.find_element(By.XPATH,"..//select")
                if fieldTypeLocator:
                    select = Select(label.find_element(By.XPATH,"..//select"))
                    selected_option = select.first_selected_option
                    getFieldValue = selected_option.text
                    if (getFieldValue == 'Select an option'):
                        fieldSatisfied = False
                    else:
                        fieldSatisfied = True
            
            # logging.info(f"Got Pre-filled Field Value (if any):{getFieldValue}, Field already satisfied: {fieldSatisfied}")
            # Exceptional Case
            if ('Dates of employment' in actualLabelText):
                fieldSatisfied = False
            
            metaInfo[actualLabelText]['fieldSatisfied'] =fieldSatisfied
            metaInfo[actualLabelText]['pre-fieldType'] =fieldType
            metaInfo[actualLabelText]['pre-getFieldValue'] =getFieldValue

            if not fieldSatisfied:
                if len(label.find_elements(By.XPATH,"../../fieldset")):
                    fieldType = label.find_element(By.XPATH,"../../fieldset").get_attribute('aria-describedby')
                elif len(label.find_elements(By.XPATH,"../../div")):
                    fieldType = label.find_element(By.XPATH,"../../div").get_attribute('aria-describedby')
                
                # Various Field Locators
                radio_t = len(label.find_elements(By.XPATH,"../..//fieldset[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//..//input[contains(@type,'radio')]"))
                radio_f = len(label.find_elements(By.XPATH,"../..//fieldset[@aria-invalid='false']//p[@data-test-form-element-error-message='true']//..//input[contains(@type,'radio')]"))
                sl_t = len(label.find_elements(By.XPATH,"../..//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//..//input[@type='text']"))
                sl_f = len(label.find_elements(By.XPATH,"../..//div[@aria-invalid='false']//p[@data-test-form-element-error-message='true']//..//input[@type='text']"))
                dd_t = len(label.find_elements(By.XPATH,"../..//div[@aria-invalid='true']//p[@data-test-form-element-error-message='true']//..//select"))
                dd_f = len(label.find_elements(By.XPATH,"../..//div[@aria-invalid='false']//p[@data-test-form-element-error-message='true']//..//select"))

                # Check if question with labelTxt exists else submit in questions.json
                value = self.checkKey(quesObj,actualLabelText)
                # logging.info(f"Got Answer For Client provided is:{value}")
                metaInfo[actualLabelText]['provided-answer'] =value
                # Condition Matching
                if (radio_t or radio_f):
                    # logging.info(f"Field Type: RADIO")
                    metaInfo[actualLabelText]['post-type'] ='RADIO'
                    if value and value != '_NA_':
                        # Recent Update... getting valu eas 1 or 0
                        if value == 'Yes':
                            value = '1'
                        else:
                            value = '0'
                        
                        # Check radio button type
                        if radio_t:
                            (label.find_element(By.XPATH,"../..//fieldset[@aria-invalid='true']//input[contains(@type,'radio')]//../input[@value='"+value+"']//..//label")).click()
                        if radio_f:
                            (label.find_element(By.XPATH,"../..//fieldset[@aria-invalid='false']//input[contains(@type,'radio')]//../input[@value='"+value+"']//..//label")).click()
                    else:
                        logging.error('No Value Received For Type radio')
                        if (not value or value == '_NA_') and isReqd:
                            radioDone = False
                            continue
                elif ('Dates of employment' in actualLabelText):
                    metaInfo[actualLabelText]['post-type'] ='DOE'
                    if isinstance(value,dict):
                        metaInfo[actualLabelText]['woe_provided'] =True
                        if value.get('currently_work_here'):
                            # logging.info('          user currently works here.')
                            metaInfo[actualLabelText]['currently_work_here'] =True
                            checked = self.browser.find_element_by_xpath("//input[contains(@id,'date-range-current-checkbox')]").get_attribute('checked')
                            if not checked:
                                self.browser.find_element_by_xpath("//input[contains(@id,'date-range-current-checkbox')]//../label").click()
                            # When User has providede that he currently works here, only start date element will be present and need to submit
                            # the current job start_month and start_year
                            sm_ = Select(label.find_element(By.XPATH,"//../div/div/fieldset[1]/fieldset//select[@name='startMonth']"))
                            sy_ = Select(label.find_element(By.XPATH,"//../div/div/fieldset[1]/fieldset//select[@name='startYear']"))
                            st_month = value.get('start_month')
                            st_year = value.get('start_year')
                            sm_.select_by_visible_text(st_month)
                            sy_.select_by_visible_text(st_year)
                        else:
                            # logging.info('          user provided no current status')
                            metaInfo[actualLabelText]['currently_work_here'] =False
                            sm_ = Select(element.find_element(By.XPATH,"//../div/div/fieldset[1]/fieldset//select[@name='startMonth']"))
                            sy_ = Select(element.find_element(By.XPATH,"//../div/div/fieldset[1]/fieldset//select[@name='startYear']"))
                            em_ = Select(element.find_element(By.XPATH,"//../div/div/fieldset/fieldset//select[@name='endMonth']"))
                            ey_ = Select(element.find_element(By.XPATH,"//../div/div/fieldset/fieldset//select[@name='endYear']"))

                            st_month = value.get('start_month')
                            st_year = value.get('start_year')
                            et_month = value.get('end_month')
                            et_year = value.get('end_year')

                            sm_.select_by_visible_text(st_month)
                            sy_.select_by_visible_text(st_year)

                            em_.select_by_visible_text(et_month)
                            ey_.select_by_visible_text(et_year)
                    else:
                        # print('Work Expreience element not handled')
                        metaInfo[actualLabelText]['woe_provided'] =False
                elif ('City' in actualLabelText):
                    metaInfo[actualLabelText]['post-type'] ='CITY'
                    try:
                        label.find_element(By.XPATH,"..//input[@type='text']").send_keys(value)
                        time.sleep(2)
                        label.find_element(By.XPATH,"..//input[@type='text']").send_keys(Keys.ARROW_DOWN,Keys.ENTER,Keys.TAB)
                    except Exception as e:
                        logging.exception('Exception with sending city, sending default',str(e))
                        label.find_element(By.XPATH,"..//input[@type='text']").send_keys('New')
                        time.sleep(2)
                        label.find_element(By.XPATH,"..//input[@type='text']").send_keys(Keys.ARROW_DOWN,Keys.ENTER,Keys.TAB)
                elif (sl_t or sl_f):
                    metaInfo[actualLabelText]['post-type'] ='SL'
                    # logging.info(f"Field Type is Single Line Text Input")
                    if value and value != '_NA_':
                        for char in value:
                            label.find_element(By.XPATH,"..//input[@type='text']").send_keys(char)
                            randomSec = random.uniform(0.0051, 0.01)
                            time.sleep(randomSec)

                    else:
                        # logging.error('No Value Received For Type Single Line Input(SL_)')
                        if (not value or value == '_NA_') and isReqd:
                            slDone = False
                            continue
                elif (dd_t or dd_f):
                    metaInfo[actualLabelText]['post-type'] ='DD'
                    # logging.info('          dd_')
                    if isinstance(value,list):
                        value = value[0]
                        # logging.info('Field is DropDown, Selecting index 0 value from list of possible values')
                    if value and value != '_NA_': 
                        s_ = Select(label.find_element(By.XPATH,"..//select"))
                        s_.select_by_visible_text(value)
                    else:
                        if (not value or value == '_NA_') and isReqd:
                            slDone = False
                            continue
            else:
                logging.info('\t\tField Value is already Set. Skipping')
                pass

        metaInfo['slDone'] = slDone
        metaInfo['radioDone']=radioDone
        logging.info('\t\t__________EXIT_______\t\t')
        return metaInfo
    
    def Merge(self,dict1, dict2): 
        res = {**dict1, **dict2} 
        return res

    def load_page(self, sleep=1):
        scroll_page = 0
        while scroll_page < 4000:
            self.browser.execute_script("window.scrollTo(0,"+str(scroll_page)+" );")
            scroll_page += 200
            time.sleep(sleep)

        if sleep != 1:
            self.browser.execute_script("window.scrollTo(0,0);")
            time.sleep(sleep * 3)

        page = BeautifulSoup(self.browser.page_source, "lxml")
        return page

    def avoid_lock(self):
        x, _ = pyautogui.position()
        pyautogui.moveTo(x + 200, pyautogui.position().y, duration=1.0)
        pyautogui.moveTo(x, pyautogui.position().y, duration=0.5)
        pyautogui.keyDown('ctrl')
        pyautogui.press('esc')
        pyautogui.keyUp('ctrl')
        time.sleep(0.5)
        pyautogui.press('esc')

    def next_jobs_page(self, position, location, jobs_per_page):
        self.browser.get(
            "https://www.linkedin.com/jobs/search/?f_LF=f_AL&keywords=" +
            position + location + "&start="+str(jobs_per_page))
        self.avoid_lock()
        self.load_page()
        jobListPageJobIds = None
        jobListPageJobIds = self.extractJobIds()
        return (self.browser, jobs_per_page,jobListPageJobIds)
    
    def next_jobs_page_custom_uri(self, customUri,jobs_per_page):
        self.browser.get(customUri+"&start="+str(jobs_per_page))
        self.avoid_lock()
        self.load_page()
        current_url = self.browser.current_url
        logging.info(f"Next URL IS: {current_url}")
        jobListPageJobIds = None
        jobListPageJobIds = self.extractJobIds()
        return (self.browser, jobs_per_page,jobListPageJobIds)
    
    def custom_jobs_page(self, customUri):
        logging.info(f'\tCustom Base (0) Url Is: {customUri}')
        self.browser.get(customUri)
        self.avoid_lock()
        self.load_page()
        jobListPageJobIds = self.extractJobIds()
        return (self.browser,jobListPageJobIds)

    def finish_apply(self):
        self.browser.close()

    def create_resume(self,metadata):
        try:
            logging.info('\t *Entered Docx & PDF Creation SubProcess. Ensure WORD is installed on your system.')
            import datetime
            from docx import Document
            import comtypes.client
            wdFormatPDF = 17

            TODAY = datetime.datetime.now()
            TODAY = TODAY.strftime("%d-%B-%Y")

            # Create DIrectory related to current day
            path_to_current = os.path.join(wkdir, 'resume\\'+str(self.TODAY))
            if os.path.isdir(path_to_current) is False:
                os.mkdir(path_to_current)
            
            # Create New DOCX FILE
            docFileName = wkdir+'\\resume\\'+str(self.TODAY)+'\\'+metadata['file_name']+'.docx'
            if not (os.path.isfile(docFileName)):
                logging.info('\t *Docx doesnt existed. Creating now')
                # Path to master Template
                doc = Document('resume\\master_template.docx')
                
                # Writable Parts to NEW DOCS
                JOB_TITLE=metadata.get('job_title') if metadata.get('job_title') != None else '-'
                APLD_FOR_JOB=TODAY
                COMPANY=metadata.get('job_company') if metadata.get('job_company') != None else '-'
                LOCATION=metadata.get('job_location') if metadata.get('job_location') != None else '-'
                DESCRIPTION=metadata.get('job_description') if metadata.get('job_description') != None else '-'
                COUNTRY = metadata.get('job_country') if metadata.get('job_country') != None else '-'

                doc.add_page_break()
                doc.add_paragraph('JOB_TITLE:'+JOB_TITLE+'\n\n'+'APLD_FOR_JOB:'+APLD_FOR_JOB+'\n\n'+'COUNTRY:'+COUNTRY+'\n\n'+'COMPANY:'+COMPANY+'\n\n'+'LOCATION:'+LOCATION+'\n\n'+'JOB DESCRIPTION:'+DESCRIPTION)
                
                for para in doc.paragraphs:
                    if 'PH-JOB-TITLE' in para.text:
                        inline = para.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if 'PH-JOB-TITLE' in inline[i].text:
                                text = inline[i].text.replace('PH-JOB-TITLE', JOB_TITLE)
                                inline[i].text = text
                    
                    if 'PH-FULL-NAME' in para.text:
                        inline = para.runs
                        # Loop added to work with runs (strings with same style)
                        for i in range(len(inline)):
                            if 'PH-FULL-NAME' in inline[i].text:
                                text = inline[i].text.replace('PH-FULL-NAME', self.full_name)
                                inline[i].text = text
                
                # Create New DOCX FILE
                    # os.remove(docFileName)
                    # time.sleep(1)
                doc.save(docFileName)
                time.sleep(0.5)
            else:
                logging.info('\t *Docx With current JOBID existed.')
            
            # Convert docx to pdf
            iFile = docFileName
            oFile = os.path.join(wkdir, 'resume\\'+str(self.TODAY)+'\\'+metadata['file_name']+'.pdf')

            if not (os.path.isfile(oFile)):
                logging.info('\t *PDF missing. Creating One.')
                word = comtypes.client.CreateObject('Word.Application')
                doc_ = word.Documents.Open(iFile)
                # os.remove(oFile)
                # time.sleep(1)
                time.sleep(1)
                doc_.SaveAs(oFile, FileFormat=wdFormatPDF)
                doc_.Close()
                word.Quit()
                time.sleep(1)
            else:
                logging.info('\t *PDF for current JOBID existed. Skipping')

        except Exception as createExcept:
            logging.exception(str(createExcept))

    def writeToJson(self,dict_,key,value):
        key = key.split('?')
        key = key[0]
        dict_[key]=value
        with open('questions.json', 'w') as json_file:
            json.dump(dict_, json_file)

    def checkKey(self,dict_, text_):
        value_ = ''
        for key,value in dict_.items():
            if key in text_ or key == text_:
                value_= dict_[key]
                break
        if not value_:
            self.writeToJson(dict_,text_,'_NA_')
        return value_
    
    def extractJobIds(self):
        logging.info('@@@@@@@@@@@@ Scrapping Page Results for Job Id"s @@@@@@@@@@')
        IDs = []
        for i in range(1,26):
            jobID = None
            elem= None
            xpath= None
            temp= None
            xpath = "(//li[contains(@class, 'occludable-update')])[{i}]".format(i=i)
            elem = self.browser.find_element_by_xpath(xpath)
            if elem:
                self.browser.execute_script("arguments[0].scrollIntoView();", elem)
                time.sleep(0.5)
                # get job links
                link = None
                xpath_div = "(//li[contains(@class, 'occludable-update')])[{i}]//div[@data-job-id]".format(i=i)
                link = self.browser.find_element_by_xpath(xpath_div)
                if link:
                    # get job ID of each job link
                    temp = link.get_attribute("data-job-id")
                    jobID = temp.split(":")[-1]
                    IDs.append(int(jobID))
                else:
                    logging.warning(' * No List Element Found ! * ')
        IDs = set(IDs)
        logging.info("Current Page Jobs Length: {jobsL}".format(jobsL=len(IDs)))
        return IDs
    
    def writeError(self,message):
        try:
            import datetime
            errFolderPath = os.path.join(wkdir, 'error')
            if os.path.isdir(errFolderPath) is False:
                os.mkdir(errFolderPath)
            errorFileName = wkdir+'\\error\\'+str(self.TODAY)+'_error'+'.txt'
            f = open(errorFileName, "a")
            f.write(str(message))
        except Exception as e:
            print(str(e))
            pass